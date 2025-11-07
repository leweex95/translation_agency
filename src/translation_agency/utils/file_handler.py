"""Document handling utilities for various file formats with formatting preservation."""

import os
from pathlib import Path
from typing import Tuple, Union, Dict, List, Any, Optional
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import PyPDF2
import copy


class FormattedRun:
    """Represents a formatted text segment with styling information."""
    
    def __init__(self, text: str, run_obj: Optional[Any] = None):
        """Initialize a formatted run."""
        self.text = text
        self.font_name = None
        self.font_size = None
        self.bold = False
        self.italic = False
        self.underline = False
        self.color_rgb = None
        self.highlight_color = None
        
        # Extract formatting from docx run object
        if run_obj is not None:
            self._extract_formatting(run_obj)
    
    def _extract_formatting(self, run_obj):
        """Extract formatting properties from docx run object."""
        font = run_obj.font
        
        if font.name:
            self.font_name = font.name
        if font.size:
            self.font_size = font.size
        if font.bold:
            self.bold = True
        if font.italic:
            self.italic = True
        if font.underline:
            self.underline = True
        if font.color and font.color.rgb:
            self.color_rgb = font.color.rgb
        
        # Extract highlight color from run element
        try:
            highlight = run_obj._element.find(qn('w:highlight'))
            if highlight is not None:
                self.highlight_color = highlight.get(qn('w:val'))
        except:
            pass
    
    def apply_to_run(self, run_obj):
        """Apply stored formatting to a docx run object."""
        font = run_obj.font
        
        if self.font_name:
            font.name = self.font_name
        if self.font_size:
            font.size = self.font_size
        if self.bold:
            font.bold = True
        if self.italic:
            font.italic = True
        if self.underline:
            font.underline = True
        if self.color_rgb:
            font.color.rgb = self.color_rgb
        if self.highlight_color:
            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), self.highlight_color)
            run_obj._element.append(highlight)


class FormattedParagraph:
    """Represents a paragraph with its structure and formatting."""
    
    def __init__(self, para_obj: Any = None):
        """Initialize a formatted paragraph."""
        self.runs: List[FormattedRun] = []
        self.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self.style_name = None
        self.left_indent = None
        self.right_indent = None
        self.space_before = None
        self.space_after = None
        self.line_spacing = None
        self.background_color = None
        self.level = 0  # For outline level (lists, etc.)
        
        if para_obj is not None:
            self._extract_from_paragraph(para_obj)
    
    def _extract_from_paragraph(self, para_obj):
        """Extract formatting from a docx paragraph."""
        # Extract text runs
        for run in para_obj.runs:
            formatted_run = FormattedRun(run.text, run)
            self.runs.append(formatted_run)
        
        # Extract paragraph formatting
        if para_obj.alignment is not None:
            self.alignment = para_obj.alignment
        
        if para_obj.style:
            self.style_name = para_obj.style.name
        
        pf = para_obj.paragraph_format
        if pf.left_indent:
            self.left_indent = pf.left_indent
        if pf.right_indent:
            self.right_indent = pf.right_indent
        if pf.space_before:
            self.space_before = pf.space_before
        if pf.space_after:
            self.space_after = pf.space_after
        if pf.line_spacing:
            self.line_spacing = pf.line_spacing
        
        # Extract outline level (for lists)
        try:
            pPr = para_obj._element.get_or_add_pPr()
            ilvl = pPr.find(qn('w:ilvl'))
            if ilvl is not None:
                self.level = int(ilvl.get(qn('w:val')))
        except:
            pass
        
        # Extract background color (shading)
        try:
            pPr = para_obj._element.get_or_add_pPr()
            shd = pPr.find(qn('w:shd'))
            if shd is not None:
                self.background_color = shd.get(qn('w:fill'))
        except:
            pass
    
    def get_text(self) -> str:
        """Get the text content of this paragraph."""
        return ''.join(run.text for run in self.runs)
    
    def set_text(self, text: str) -> None:
        """Set the text content and clear existing runs."""
        if not self.runs:
            self.runs = [FormattedRun(text)]
        else:
            # Update first run and keep formatting
            self.runs[0].text = text
            # Remove other runs
            self.runs = self.runs[:1]
    
    def recreate_in_document(self, doc: Document, text: str) -> Any:
        """Recreate paragraph in document with formatting applied."""
        # Determine style
        style = self.style_name if self.style_name else None
        para = doc.add_paragraph(style=style)
        
        # Apply paragraph formatting
        pf = para.paragraph_format
        if self.left_indent:
            pf.left_indent = self.left_indent
        if self.right_indent:
            pf.right_indent = self.right_indent
        if self.space_before:
            pf.space_before = self.space_before
        if self.space_after:
            pf.space_after = self.space_after
        if self.line_spacing:
            pf.line_spacing = self.line_spacing
        
        para.alignment = self.alignment
        
        # Apply background color
        if self.background_color:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), self.background_color)
            para._element.get_or_add_pPr().append(shading_elm)
        
        # Add the entire translated text as a single run
        # This preserves the paragraph-level structure while updating content
        run = para.add_run(text)
        
        # Apply formatting from first run (if available)
        # This maintains the overall paragraph formatting style
        if self.runs:
            self.runs[0].apply_to_run(run)
        
        return para


class FormattedTable:
    """Represents a table with its structure and formatting."""
    
    def __init__(self, table_obj: Any = None):
        """Initialize a formatted table."""
        self.rows: List[List[str]] = []
        self.table_style = None
        self.cell_formatting: List[List[Dict]] = []
        self.header_row = False
        
        if table_obj is not None:
            self._extract_from_table(table_obj)
    
    def _extract_from_table(self, table_obj):
        """Extract content and formatting from docx table."""
        self.table_style = table_obj.style.name if table_obj.style else 'Table Grid'
        
        for row_idx, row in enumerate(table_obj.rows):
            row_data = []
            row_formatting = []
            
            for cell in row.cells:
                # Extract cell text (combine all paragraphs)
                cell_text = '\n'.join(
                    p.text for p in cell.paragraphs if p.text.strip()
                )
                row_data.append(cell_text)
                
                # Extract cell formatting (color, shading, etc.)
                cell_format = {}
                try:
                    # Check for background color
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        cell_format['background_color'] = shd.get(qn('w:fill'))
                except:
                    pass
                
                row_formatting.append(cell_format)
            
            self.rows.append(row_data)
            self.cell_formatting.append(row_formatting)
    
    def get_text_rows(self) -> List[str]:
        """Get table content as text (rows separated by newline)."""
        result = []
        for row in self.rows:
            result.append(' | '.join(row))
        return result
    
    def recreate_in_document(self, doc: Document, translated_rows: List[List[str]]) -> Any:
        """Recreate table in document with formatting applied."""
        if not self.rows:
            return None
        
        num_rows = len(self.rows)
        num_cols = len(self.rows[0]) if self.rows else 0
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        if self.table_style:
            table.style = self.table_style
        
        # Fill table with translated content and apply formatting
        for row_idx in range(num_rows):
            for col_idx in range(num_cols):
                cell = table.rows[row_idx].cells[col_idx]
                
                # Get translated text or original if not available
                if row_idx < len(translated_rows) and col_idx < len(translated_rows[row_idx]):
                    text = translated_rows[row_idx][col_idx]
                else:
                    text = self.rows[row_idx][col_idx] if col_idx < len(self.rows[row_idx]) else ''
                
                cell.text = text
                
                # Apply cell formatting
                if row_idx < len(self.cell_formatting) and col_idx < len(self.cell_formatting[row_idx]):
                    cell_format = self.cell_formatting[row_idx][col_idx]
                    if 'background_color' in cell_format:
                        tcPr = cell._tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), cell_format['background_color'])
                        tcPr.append(shd)
        
        return table


class DocumentStructure:
    """Represents the full structure of a document."""
    
    def __init__(self):
        """Initialize document structure."""
        self.paragraphs: List[Union[FormattedParagraph, FormattedTable]] = []
    
    def add_paragraph(self, para: FormattedParagraph) -> None:
        """Add a paragraph to the structure."""
        self.paragraphs.append(para)
    
    def add_table(self, table: FormattedTable) -> None:
        """Add a table to the structure."""
        self.paragraphs.append(table)
    
    def get_all_text(self) -> str:
        """Extract all text from structure."""
        texts = []
        for item in self.paragraphs:
            if isinstance(item, FormattedParagraph):
                if item.get_text().strip():
                    texts.append(item.get_text())
            elif isinstance(item, FormattedTable):
                texts.extend(item.get_text_rows())
        return '\n\n'.join(texts)


class DocumentHandler:
    """Handle reading and writing of different document formats."""
    
    SUPPORTED_FORMATS = {'.txt', '.md', '.markdown', '.docx', '.pdf'}
    
    @classmethod
    def read_document_with_formatting(cls, file_path: Union[str, Path]) -> Tuple[DocumentStructure, str]:
        """
        Read document and preserve its structure and formatting.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (DocumentStructure, file_extension)
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = path.suffix.lower()
        
        if file_ext not in cls.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {cls.SUPPORTED_FORMATS}")
        
        if file_ext == '.docx':
            return cls._read_docx_with_formatting(path), file_ext
        else:
            # Fall back to plain text reading for other formats
            content = cls.read_document(file_path)[0]
            structure = DocumentStructure()
            structure.add_paragraph(FormattedParagraph())
            structure.paragraphs[0].set_text(content)
            return structure, file_ext
    
    @classmethod
    def write_document_with_formatting(cls, structure: DocumentStructure, translated_texts: List[str],
                                      output_path: Union[str, Path], original_format: str = '.txt') -> Path:
        """
        Write document preserving structure and formatting from the original.
        
        Args:
            structure: DocumentStructure object containing original formatting
            output_path: Path where to save the file
            translated_texts: List of translated text segments in order
            original_format: Original file format
            
        Returns:
            The path where the file was written
        """
        path = Path(output_path)
        path.parent.mkdir(parents=True, exist_ok=True)
        
        if original_format == '.docx':
            return cls._write_docx_with_formatting(structure, translated_texts, path)
        else:
            # For other formats, combine texts and write as basic format
            content = '\n\n'.join(translated_texts)
            return cls.write_document(content, path, original_format)
    
    @classmethod
    def read_document(cls, file_path: Union[str, Path]) -> Tuple[str, str]:
        """
        Read document content and return (content, format).
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Tuple of (content, file_extension)
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_ext = path.suffix.lower()
        
        if file_ext not in cls.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported: {cls.SUPPORTED_FORMATS}")
        
        if file_ext == '.txt':
            return cls._read_txt(path), file_ext
        elif file_ext in ['.md', '.markdown']:
            return cls._read_markdown(path), file_ext
        elif file_ext == '.docx':
            return cls._read_docx(path), file_ext
        elif file_ext == '.pdf':
            return cls._read_pdf(path), file_ext
    
    @classmethod
    def write_document(cls, content: str, output_path: Union[str, Path], 
                      original_format: str = '.txt') -> Path:
        """
        Write content to document in specified format.
        
        Args:
            content: Text content to write
            output_path: Path where to save the file
            original_format: Original file format to preserve structure
            
        Returns:
            The actual path where the file was written
        """
        path = Path(output_path)
        
        # Ensure output directory exists
        path.parent.mkdir(parents=True, exist_ok=True)
        
        if original_format == '.txt':
            cls._write_txt(content, path)
            return path
        elif original_format in ['.md', '.markdown']:
            cls._write_markdown(content, path)
            return path
        elif original_format == '.docx':
            cls._write_docx(content, path)
            return path
        elif original_format == '.pdf':
            # PDF input -> text output (can't write back to PDF easily)
            actual_path = path.with_suffix('.txt')
            cls._write_txt(content, actual_path)
            return actual_path
        else:
            raise ValueError(f"Cannot write to unsupported format: {original_format}")
    
    @staticmethod
    def _read_docx_with_formatting(path: Path) -> DocumentStructure:
        """Read DOCX file and preserve all formatting."""
        doc = Document(path)
        structure = DocumentStructure()
        
        # Simple approach: process all paragraphs and tables in document order
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                formatted_para = FormattedParagraph(para)
                structure.add_paragraph(formatted_para)
        
        for table in doc.tables:
            formatted_table = FormattedTable(table)
            structure.add_table(formatted_table)
        
        return structure
    
    @staticmethod
    def _write_docx_with_formatting(structure: DocumentStructure, translated_texts: List[str],
                                    path: Path) -> Path:
        """Write DOCX file preserving formatting from original structure."""
        # Ensure .docx extension
        if path.suffix.lower() != '.docx':
            path = path.with_suffix('.docx')
        
        doc = Document()
        text_index = 0
        
        for item in structure.paragraphs:
            if isinstance(item, FormattedParagraph):
                # Get translated text for this paragraph
                if text_index < len(translated_texts):
                    translated = translated_texts[text_index]
                    text_index += 1
                else:
                    translated = item.get_text()
                
                # Skip empty paragraphs (but preserve in structure)
                if not translated.strip():
                    item.recreate_in_document(doc, '')
                else:
                    item.recreate_in_document(doc, translated)
            
            elif isinstance(item, FormattedTable):
                # For tables, we need to split text into rows
                table_translated_rows = []
                for row_idx in range(len(item.rows)):
                    row_cells = []
                    for col_idx in range(len(item.rows[row_idx])):
                        if text_index < len(translated_texts):
                            row_cells.append(translated_texts[text_index])
                            text_index += 1
                        else:
                            row_cells.append(item.rows[row_idx][col_idx])
                    table_translated_rows.append(row_cells)
                
                item.recreate_in_document(doc, table_translated_rows)
        
        doc.save(path)
        return path
    
    @staticmethod
    def _read_txt(path: Path) -> str:
        """Read plain text file."""
        try:
            return path.read_text(encoding='utf-8')
        except UnicodeDecodeError as e:
            # If UTF-8 decoding fails, show the file content for debugging
            print(f"[ERROR] UTF-8 decoding failed for file: {path}")
            print("[CONTENT] File content (first 500 bytes):")
            try:
                with open(path, 'rb') as f:
                    content = f.read(500)
                    # Safely print binary content
                    try:
                        print(content.decode('utf-8', errors='replace'))
                    except UnicodeEncodeError:
                        print(f"[BINARY] Binary content detected ({len(content)} bytes)")
                    if len(content) == 500:
                        print("... (truncated)")
            except Exception as read_error:
                print(f"[ERROR] Could not read file content: {read_error}")
            raise ValueError(f"File {path} is not valid UTF-8 text. This appears to be a binary file or uses a different encoding. Supported formats: {DocumentHandler.SUPPORTED_FORMATS}") from e
    
    @staticmethod
    def _read_markdown(path: Path) -> str:
        """Read markdown file (return as plain text for translation)."""
        try:
            return path.read_text(encoding='utf-8')
        except UnicodeDecodeError as e:
            # If UTF-8 decoding fails, show the file content for debugging
            print(f"[ERROR] UTF-8 decoding failed for file: {path}")
            print("[CONTENT] File content (first 500 bytes):")
            try:
                with open(path, 'rb') as f:
                    content = f.read(500)
                    # Safely print binary content
                    try:
                        print(content.decode('utf-8', errors='replace'))
                    except UnicodeEncodeError:
                        print(f"[BINARY] Binary content detected ({len(content)} bytes)")
                    if len(content) == 500:
                        print("... (truncated)")
            except Exception as read_error:
                print(f"[ERROR] Could not read file content: {read_error}")
            raise ValueError(f"File {path} is not valid UTF-8 text. This appears to be a binary file or uses a different encoding. Supported formats: {DocumentHandler.SUPPORTED_FORMATS}") from e
    
    @staticmethod
    def _read_docx(path: Path) -> str:
        """Read DOCX file and extract text content."""
        doc = Document(path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    content.append(' | '.join(row_text))
        
        return '\n\n'.join(content)
    
    @staticmethod
    def _read_pdf(path: Path) -> str:
        """Read PDF file and extract text content."""
        content = []
        
        with open(path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    content.append(page_text.strip())
        
        if not content:
            raise ValueError(f"No text could be extracted from PDF: {path}")
        
        return '\n\n'.join(content)
    
    @staticmethod
    def _write_txt(content: str, path: Path) -> None:
        """Write plain text file."""
        path.write_text(content, encoding='utf-8')
    
    @staticmethod
    def _write_markdown(content: str, path: Path) -> None:
        """Write markdown file."""
        # Ensure .md extension
        if path.suffix.lower() not in ['.md', '.markdown']:
            path = path.with_suffix('.md')
        path.write_text(content, encoding='utf-8')
    
    @staticmethod
    def _write_docx(content: str, path: Path) -> None:
        """Write content to DOCX file."""
        # Ensure .docx extension
        if path.suffix.lower() != '.docx':
            path = path.with_suffix('.docx')
            
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = content.split('\n\n')
        
        for para_text in paragraphs:
            if para_text.strip():
                # Handle potential table-like content (contains |)
                if ' | ' in para_text:
                    # Create table for structured content
                    rows = [row.strip() for row in para_text.split('\n') if ' | ' in row]
                    if rows:
                        cols = rows[0].split(' | ')
                        table = doc.add_table(rows=len(rows), cols=len(cols))
                        table.style = 'Table Grid'
                        
                        for i, row_text in enumerate(rows):
                            cells = row_text.split(' | ')
                            for j, cell_text in enumerate(cells):
                                if j < len(table.rows[i].cells):
                                    table.rows[i].cells[j].text = cell_text.strip()
                else:
                    # Regular paragraph
                    doc.add_paragraph(para_text.strip())
        
        doc.save(path)
    
    @classmethod
    def get_output_filename(cls, input_path: Union[str, Path], step_name: str, 
                           output_dir: Union[str, Path], timestamp: Optional[str] = None) -> Path:
        """
        Generate output filename for a pipeline step.
        
        Args:
            input_path: Original input file path
            step_name: Name of the pipeline step
            output_dir: Output directory
            timestamp: Optional timestamp string (YYYYMMDD_HHMMSS format). If None, generates new timestamp.
            
        Returns:
            Path object for the output file
        """
        from datetime import datetime
        
        input_path = Path(input_path)
        output_dir = Path(output_dir)
        
        # Use provided timestamp or generate new one
        if timestamp is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Create document-specific folder with datetime stamp: output_dir/YYYYMMDD_HHMMSS_document_name/
        document_name = input_path.stem  # filename without extension
        document_dir = output_dir / f"{timestamp}_{document_name}"
        document_dir.mkdir(parents=True, exist_ok=True)
        
        # Create filename: step_name.ext (inside document folder)
        ext = input_path.suffix
        
        output_filename = f"{step_name}{ext}"
        return document_dir / output_filename